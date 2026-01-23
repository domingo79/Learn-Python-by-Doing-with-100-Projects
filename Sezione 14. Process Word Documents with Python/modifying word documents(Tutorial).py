from docx import Document
import os
"""
document = Document()
document.add_paragraph("It was a dark and stormy night.")
document.save("dark-and-stormy.docx")
document = Document("dark-and-stormy.docx")
document.paragraphs[0].text
"""
FOLDER = os.path.dirname(os.path.abspath(__file__))
RESOURCES = os.path.join(FOLDER, 'resources')
OUTPUT = os.path.join(FOLDER, 'output')
# percorso dei singoli file
doc_path_1, doc_path_2 = [os.path.join(
    RESOURCES, f) for f in os.listdir(RESOURCES) if f.startswith("panda")]

# selezioniamo i due documenti
doc_1 = Document(doc_path_1)
doc_2 = Document(doc_path_2)

parag = doc_2.paragraphs[0]

paragraphs = doc_1.paragraphs
paragraphs[1]._element.addnext(parag._element)
doc_1.save(os.path.join(OUTPUT, 'panda.docx'))
