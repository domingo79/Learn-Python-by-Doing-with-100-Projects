from docx import Document
from copy import deepcopy
import os

"""Creare un programma che unisca i documenti di Microsoft Word.
(1) Quando l'utente esegue il programma, il programma dovrebbe cercare i file .docx 
dalla cartella resources. 
(2) Il programma dovrebbe salvare il contenuto di tutti i documenti in un unico (nuovo) documento 
e salvare il nuovo documento nella cartella output con il nome di word_documents.
"""
FOLDER = os.path.dirname(os.path.abspath(__file__))
RESOURCES = os.path.join(FOLDER, 'resources')
OUTPUT = os.path.join(FOLDER, 'output')

file_paths = [os.path.join(RESOURCES, dir_name) for dir_name in os.listdir(
    RESOURCES) if not (dir_name.startswith("panda") or dir_name.startswith("."))]

# Crea un nuovo oggetto Documento per il merge del contenuto dei file
merged_document = Document()

# Itera su tutti i file_paths nella cartella specificata
for filepath in file_paths:
    doc = Document(filepath)

    # Aggiungi ogni paragrafo del documento corrente e al documento merge
    for para in doc.paragraphs:
        new_para = merged_document.add_paragraph()
        new_para._element.addprevious(para._element)

# Salvataggio
merged_document.save(os.path.join(OUTPUT, 'word_documents.docx'))
